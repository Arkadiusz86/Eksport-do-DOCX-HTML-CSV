import sys
import os
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if len(sys.argv) != 2:
    print("Użycie: python parse_nmap.py <plik_xml>")
    sys.exit(1)

xml_file = sys.argv[1]
docx_file = os.path.splitext(xml_file)[0] + '.docx'

tree = ET.parse(xml_file)
root = tree.getroot()

# Metadane
scan_args  = root.attrib.get('args', 'brak')
scan_start = root.attrib.get('startstr', 'brak')
scan_ver   = root.attrib.get('version', 'brak')

runstats = root.find('runstats')
elapsed, hosts_up, hosts_down, hosts_total = '?', '?', '?', '?'
if runstats is not None:
    fin = runstats.find('finished')
    if fin is not None: elapsed = fin.attrib.get('elapsed', '?') + 's'
    he = runstats.find('hosts')
    if he is not None:
        hosts_up    = he.attrib.get('up', '?')
        hosts_down  = he.attrib.get('down', '?')
        hosts_total = he.attrib.get('total', '?')

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_bold_para(doc, text, size=11, color=None, space_before=6, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_kv(doc, key, value):
    """Dodaj linię klucz: wartość"""
    if not value:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    k = p.add_run(key + ': ')
    k.bold = True
    k.font.size = Pt(10)
    v = p.add_run(value)
    v.font.size = Pt(10)

def make_header_row(table, labels, bg='2C3E50'):
    row = table.rows[0]
    for i, label in enumerate(labels):
        if i < len(row.cells):
            cell = row.cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(cell, bg)

# ── Dokument ──────────────────────────────────────────────────────────────────

doc = Document()

# Marginesy
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# Tytuł
title = doc.add_heading('Raport Nmap', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadane skanowania
add_kv(doc, 'Data skanowania', scan_start)
add_kv(doc, 'Wersja Nmap', scan_ver)
add_kv(doc, 'Czas skanowania', elapsed)
add_kv(doc, 'Komenda', scan_args)

# Podsumowanie
doc.add_paragraph()
add_bold_para(doc, 'Podsumowanie', size=12)
sum_table = doc.add_table(rows=2, cols=3)
sum_table.style = 'Table Grid'
labels_sum = ['Hosty online', 'Hosty offline', 'Łącznie']
values_sum = [hosts_up, hosts_down, hosts_total]
colors_sum = ['27AE60', 'E74C3C', '2980B9']
for i, (l, v, c) in enumerate(zip(labels_sum, values_sum, colors_sum)):
    hc = sum_table.rows[0].cells[i]
    vc = sum_table.rows[1].cells[i]
    hc.text = l
    hc.paragraphs[0].runs[0].bold = True
    hc.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hc, c)
    hc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    vc.text = v
    vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    vc.paragraphs[0].runs[0].bold = True
    vc.paragraphs[0].runs[0].font.size = Pt(14)

doc.add_paragraph()

# ── Hosty ─────────────────────────────────────────────────────────────────────

for host in root.findall('host'):
    status = host.find('status')
    host_state  = status.attrib.get('state', '') if status is not None else ''
    host_reason = status.attrib.get('reason', '') if status is not None else ''

    ip, ipv6, mac, mac_vendor = '', '', '', ''
    for addr in host.findall('address'):
        atype = addr.attrib.get('addrtype', '')
        if atype == 'ipv4':   ip = addr.attrib.get('addr', '')
        elif atype == 'ipv6': ipv6 = addr.attrib.get('addr', '')
        elif atype == 'mac':
            mac = addr.attrib.get('addr', '')
            mac_vendor = addr.attrib.get('vendor', '')

    hostnames = host.find('hostnames')
    hostname = ''
    if hostnames is not None:
        hn = hostnames.find('hostname')
        if hn is not None: hostname = hn.attrib.get('name', '')

    os_name, os_accuracy = '', ''
    os_elem = host.find('os')
    if os_elem is not None:
        osmatch = os_elem.find('osmatch')
        if osmatch is not None:
            os_name     = osmatch.attrib.get('name', '')
            os_accuracy = osmatch.attrib.get('accuracy', '') + '%'

    uptime_seconds, uptime_lastboot = '', ''
    uptime = host.find('uptime')
    if uptime is not None:
        uptime_seconds  = uptime.attrib.get('seconds', '') + 's'
        uptime_lastboot = uptime.attrib.get('lastboot', '')

    distance = ''
    dist_elem = host.find('distance')
    if dist_elem is not None:
        distance = dist_elem.attrib.get('value', '') + ' hop(s)'

    display_ip = ip or ipv6 or 'nieznany'
    hdr_color = (0x27, 0xAE, 0x60) if host_state == 'up' else (0xE7, 0x4C, 0x3C)

    p = add_bold_para(doc, f'Host: {display_ip}', size=13, color=hdr_color, space_before=12)
    p.paragraph_format.keep_with_next = True

    add_kv(doc, 'Stan', f'{host_state} ({host_reason})')
    if hostname:       add_kv(doc, 'Hostname', hostname)
    if ipv6:           add_kv(doc, 'IPv6', ipv6)
    if mac:            add_kv(doc, 'MAC', f'{mac}' + (f' ({mac_vendor})' if mac_vendor else ''))
    if os_name:        add_kv(doc, 'System operacyjny', f'{os_name} (pewność: {os_accuracy})')
    if uptime_seconds: add_kv(doc, 'Uptime', f'{uptime_seconds} (ostatni restart: {uptime_lastboot})')
    if distance:       add_kv(doc, 'Odległość sieciowa', distance)

    # Skrypty hosta
    hostscript = host.find('hostscript')
    if hostscript is not None:
        add_bold_para(doc, 'Skrypty hosta (NSE):', size=10, space_before=4, space_after=2)
        for script in hostscript.findall('script'):
            sid  = script.attrib.get('id', '')
            sout = script.attrib.get('output', '')
            p2 = doc.add_paragraph(style='List Bullet')
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after  = Pt(1)
            r1 = p2.add_run(sid + ': ')
            r1.bold = True
            r1.font.size = Pt(9)
            r2 = p2.add_run(sout)
            r2.font.size = Pt(9)

    # Tabela portów
    ports = host.find('ports')
    port_list = ports.findall('port') if ports is not None else []

    if port_list:
        add_bold_para(doc, 'Porty:', size=10, space_before=6, space_after=2)

        col_labels = ['Port/Proto', 'Stan', 'Powód', 'Usługa', 'Produkt', 'Wersja', 'Info dodatkowe', 'CPE', 'Skrypty NSE']
        col_widths = [2.2, 1.6, 2.0, 2.2, 2.8, 2.2, 2.8, 3.0, 4.0]

        tbl = doc.add_table(rows=1 + len(port_list), cols=len(col_labels))
        tbl.style = 'Table Grid'
        make_header_row(tbl, col_labels)
        set_col_widths(tbl, col_widths)

        for ri, port in enumerate(port_list, start=1):
            protocol = port.attrib.get('protocol', '')
            portid   = port.attrib.get('portid', '')

            state_elem  = port.find('state')
            port_state  = state_elem.attrib.get('state', '') if state_elem is not None else ''
            port_reason = state_elem.attrib.get('reason', '') if state_elem is not None else ''

            svc = port.find('service')
            svc_name, svc_product, svc_version, svc_extrainfo = '', '', '', ''
            cpe = ''
            if svc is not None:
                svc_name     = svc.attrib.get('name', '')
                svc_product  = svc.attrib.get('product', '')
                svc_version  = svc.attrib.get('version', '')
                svc_extrainfo= svc.attrib.get('extrainfo', '')
                cpe_list     = [c.text for c in svc.findall('cpe') if c.text]
                cpe          = ', '.join(cpe_list)

            nse_parts = []
            for script in port.findall('script'):
                sid  = script.attrib.get('id', '')
                sout = script.attrib.get('output', '')
                nse_parts.append(f'{sid}: {sout}')
            nse_text = ' | '.join(nse_parts)

            values = [
                f'{portid}/{protocol}', port_state, port_reason,
                svc_name, svc_product, svc_version, svc_extrainfo, cpe, nse_text
            ]
            row = tbl.rows[ri]
            for ci, val in enumerate(values):
                cell = row.cells[ci]
                cell.text = val or ''
                run = cell.paragraphs[0].runs
                if run:
                    run[0].font.size = Pt(8)
                # Kolor stanu portu
                if ci == 1 and val:
                    color = (0x27,0xAE,0x60) if val == 'open' else \
                            (0xE6,0x7E,0x22) if val == 'filtered' else (0x7F,0x8C,0x8D)
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(*color)
                        cell.paragraphs[0].runs[0].bold = True

        # Filtrowane porty (extraports)
        if ports is not None:
            extraports = ports.find('extraports')
            if extraports is not None:
                ep_state = extraports.attrib.get('state', '')
                ep_count = extraports.attrib.get('count', '')
                p3 = doc.add_paragraph()
                p3.paragraph_format.space_before = Pt(2)
                r = p3.add_run(f'ℹ️ {ep_count} portów w stanie: {ep_state}')
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x88,0x88,0x88)
                r.italic = True
    else:
        p_np = doc.add_paragraph()
        p_np.add_run('Brak otwartych portów.').font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

    doc.add_paragraph()  # odstęp między hostami

doc.save(docx_file)
print(f"✔️ Zapisano DOCX: {docx_file}")
